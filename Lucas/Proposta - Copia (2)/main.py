import os
import json
import time
import subprocess
from datetime import datetime
import pandas as pd
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import Select
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
import tkinter as tk
from tkinter import ttk, scrolledtext
import glob
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Dicionário com as modalidades por portal
MODALIDADES = {
    'bll': ['Selecione a modalidade', 'Pregão', 'Concorrência', 'Dispensa'],
    'comprasnet': ['Selecione a modalidade', 'Pregão', 'Concorrência', 'Dispensa'],
    'compraspublicas': ['Selecione a modalidade', 'Pregão', 'Concorrência', 'Dispensa'],
    'bnc': ['Selecione a modalidade', 'Pregão', 'Concorrência', 'Dispensa'],
    'comprasbr': ['Selecione a modalidade', 'Pregão', 'Concorrência', 'Dispensa'],
    'licitanet': ['Selecione a modalidade', 'Pregão', 'Concorrência', 'Dispensa']
}

# Mapeamento de valores das modalidades por portal
MODALIDADE_VALUES = {
    'bll': {'Selecione a modalidade': '0', 'Pregão': '1', 'Concorrência': '3', 'Dispensa': '10'},
    'comprasnet': {'Selecione a modalidade': '0', 'Pregão': '1', 'Concorrência': '3', 'Dispensa': '10'},
    'compraspublicas': {'Selecione a modalidade': '0', 'Pregão': '1', 'Concorrência': '3', 'Dispensa': '10'},
    'bnc': {'Selecione a modalidade': '0', 'Pregão': '1', 'Concorrência': '3', 'Dispensa': '10'},
    'comprasbr': {'Selecione a modalidade': '0', 'Pregão': '1', 'Concorrência': '3', 'Dispensa': '10'},
    'licitanet': {'Selecione a modalidade': '0', 'Pregão': '1', 'Concorrência': '3', 'Dispensa': '10'}
}

def fechar_popup_tour(driver, wait):
    """Função auxiliar para fechar o popup do tour se ele estiver presente"""
    try:
        # Tenta encontrar o botão de finalizar do tour
        tour_button = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "a[data-role='end']")))
        if tour_button:
            tour_button.click()
            time.sleep(1)  # Pequena pausa para garantir que o popup fechou
            print("Tour finalizado com sucesso!")
        return True
    except:
        return False

def verificar_aviso(driver, wait):
    """Função para verificar se apareceu o aviso de warning e fechá-lo"""
    try:
        # Verifica se existe o título "Aviso!" e o ícone de warning
        aviso = wait.until(EC.presence_of_element_located((By.CLASS_NAME, "swal2-warning")))
        titulo = wait.until(EC.presence_of_element_located((By.ID, "swal2-title")))
        
        if aviso.is_displayed() and titulo.is_displayed() and titulo.text == "Aviso!":
            print("Aviso detectado! Tentando fechar...")
            # Tenta encontrar e clicar no botão OK
            try:
                ok_button = wait.until(EC.element_to_be_clickable((By.CLASS_NAME, "swal2-confirm")))
                ok_button.click()
                print("Aviso fechado com sucesso!")
                time.sleep(1)  # Pequena pausa após fechar o aviso
            except Exception as e:
                print(f"Erro ao fechar aviso: {str(e)}")
            return True
    except:
        pass
    return False

class SeletorPortal:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Seletor de Portal")
        self.root.geometry("500x700")
        
        # Título
        tk.Label(
            self.root, 
            text="Selecione o Portal", 
            font=('Helvetica', 16, 'bold')
        ).pack(pady=20)
        
        # Frame para os botões com scrollbar
        container = tk.Frame(self.root)
        container.pack(fill='both', expand=True, padx=20, pady=20)
        
        # Canvas e Scrollbar
        canvas = tk.Canvas(container)
        scrollbar = tk.Scrollbar(container, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Estilo padrão para os botões
        button_style = {
            'width': 40,
            'height': 2,
            'font': ('Helvetica', 11),
            'pady': 5,
            'wraplength': 350
        }
        
        # Lista de portais
        portais = [
            ("BLL - Bolsa de Licitações do Brasil", "bll"),
            ("ComprasNet - Compras Governamentais", "comprasnet"),
            ("Portal de Compras Públicas", "compraspublicas"),
            ("BNC - Bolsa Nacional de Compras", "bnc"),
            ("Compras BR - Portal de Compras do Brasil", "comprasbr"),
            ("LicitaNet - Portal de Licitações", "licitanet")
        ]
        
        # Criando botões para cada portal
        for texto, comando in portais:
            tk.Button(
                scrollable_frame,
                text=texto,
                command=lambda cmd=comando, txt=texto: self.abrir_formulario(cmd, txt),
                **button_style
            ).pack(pady=10)
        
        # Empacotando o canvas e scrollbar
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        self.root.mainloop()
        
    def abrir_formulario(self, portal_id, portal_nome):
        # Zera o arquivo JSON antes de abrir o formulário
        with open('dados.json', 'w', encoding='utf-8') as arquivo:
            json.dump({
                "portal_id": portal_id,
                "portal_nome": portal_nome,
                "registros": []
            }, arquivo, ensure_ascii=False, indent=4)
            
        self.root.destroy()  # Fecha a janela de seleção
        
        # Abre o formulário específico do portal
        if portal_id == "bll":
            FormularioBLL(portal_id, portal_nome)
        elif portal_id == "comprasnet":
            FormularioComprasNet(portal_id, portal_nome)
        elif portal_id == "compraspublicas":
            FormularioComprasPublicas(portal_id, portal_nome)
        elif portal_id == "bnc":
            FormularioBNC(portal_id, portal_nome)
        elif portal_id == "comprasbr":
            FormularioComprasBR(portal_id, portal_nome)
        elif portal_id == "licitanet":
            FormularioLicitaNet(portal_id, portal_nome)
        else:
            messagebox.showinfo("Em Desenvolvimento", f"O formulário para o portal {portal_nome} está em desenvolvimento!")
            SeletorPortal()

class FormularioBase:
    def __init__(self, portal_id, portal_nome):
        self.root = tk.Tk()
        self.root.title(f"Formulário - {portal_nome}")
        self.root.geometry("500x700")
        
        self.portal_id = portal_id
        self.portal_nome = portal_nome
        
        # Frame para os campos
        self.campos_frame = tk.Frame(self.root)
        self.campos_frame.pack(fill='both', expand=True, padx=20, pady=20)
        
        # Frame para os botões
        self.botoes_frame = tk.Frame(self.root)
        self.botoes_frame.pack(pady=20)
        
        # Frame para a lista
        self.lista_frame = tk.Frame(self.root)
        self.lista_frame.pack(fill='both', expand=True, padx=20)
        
        # Botões
        tk.Button(self.botoes_frame, text="Adicionar", command=self.adicionar_dados).pack(side=tk.LEFT, padx=10)
        tk.Button(self.botoes_frame, text="Enviar para Chrome", command=self.enviar_chrome).pack(side=tk.LEFT, padx=10)
        tk.Button(self.botoes_frame, text="Voltar", command=self.voltar).pack(side=tk.LEFT, padx=10)
        
        # Lista de dados
        tk.Label(self.lista_frame, text="Clique em um item para removê-lo da lista:").pack(pady=5)
        self.lista_dados = tk.Listbox(self.lista_frame, width=60, height=10)
        self.lista_dados.pack(pady=10)
        self.lista_dados.bind('<<ListboxSelect>>', self.remover_item)
        
        self.carregar_dados()
    
    def criar_campo(self, frame, label_text, tipo="entry"):
        label = tk.Label(frame, text=label_text)
        label.pack(pady=5)
        
        if tipo == "entry":
            campo = tk.Entry(frame, width=50)
        elif tipo == "combobox":
            campo = ttk.Combobox(frame, width=47, state="readonly")
            campo['values'] = MODALIDADES.get(self.portal_id, [])
        
        campo.pack(pady=5)
        return campo
    
    def carregar_dados(self):
        try:
            with open('dados.json', 'r', encoding='utf-8') as arquivo:
                dados = json.load(arquivo)
                for registro in dados['registros']:
                    self.lista_dados.insert(tk.END, self.formato_lista(registro))
        except Exception as e:
            print(f"Erro ao carregar dados: {str(e)}")
    
    def formato_lista(self, registro):
        # Será sobrescrito pelas classes filhas
        pass
    
    def coletar_dados(self):
        # Será sobrescrito pelas classes filhas
        pass
    
    def adicionar_dados(self):
        novo_registro = self.coletar_dados()
        if novo_registro:
            with open('dados.json', 'r', encoding='utf-8') as arquivo:
                dados = json.load(arquivo)
            
            dados['registros'].append(novo_registro)
            
            with open('dados.json', 'w', encoding='utf-8') as arquivo:
                json.dump(dados, arquivo, ensure_ascii=False, indent=4)
            
            self.lista_dados.insert(tk.END, self.formato_lista(novo_registro))
            self.limpar_campos()
    
    def limpar_campos(self):
        # Será sobrescrito pelas classes filhas
        pass
    
    def remover_item(self, event):
        try:
            selection = self.lista_dados.curselection()
            if selection:
                index = selection[0]
                self.lista_dados.delete(index)
                
                with open('dados.json', 'r', encoding='utf-8') as arquivo:
                    dados = json.load(arquivo)
                dados['registros'].pop(index)
                with open('dados.json', 'w', encoding='utf-8') as arquivo:
                    json.dump(dados, arquivo, ensure_ascii=False, indent=4)
        except Exception as e:
            print(f"Erro ao remover item: {str(e)}")
    
    def enviar_chrome(self):
        """Envia os dados para o Chrome"""
        driver = None
        try:
            # Verifica se existem dados para processar
            if not self.lista_dados.size():
                messagebox.showwarning("Aviso", "Adicione pelo menos um registro antes de enviar para o Chrome!")
                return
            
            # Configurações do Chrome
            chrome_options = Options()
            chrome_options.add_argument("--start-maximized")
            chrome_options.add_argument("--no-sandbox")
            chrome_options.add_argument("--disable-dev-shm-usage")
            
            # Configura o diretório de download para a pasta downloads dentro do projeto
            download_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'downloads')
            if not os.path.exists(download_dir):
                os.makedirs(download_dir)
                
            chrome_options.add_experimental_option('prefs', {
                'download.default_directory': download_dir,
                'download.prompt_for_download': False,
                'download.directory_upgrade': True,
                'safebrowsing.enabled': True
            })

            try:
                # Primeira tentativa: usar o ChromeDriverManager com verificação de versão
                from webdriver_manager.chrome import ChromeDriverManager
                from selenium.webdriver.chrome.service import Service
                
                # Força o download da versão mais recente do ChromeDriver
                driver_path = ChromeDriverManager(path=os.path.dirname(os.path.abspath(__file__))).install()
                print(f"ChromeDriver instalado em: {driver_path}")
                
                service = Service(executable_path=driver_path)
                driver = webdriver.Chrome(service=service, options=chrome_options)
                print("Chrome iniciado com ChromeDriverManager")
                
            except Exception as e1:
                print(f"Primeira tentativa falhou: {str(e1)}")
                try:
                    # Segunda tentativa: usar o chromedriver local
                    local_driver = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'chromedriver.exe')
                    if os.path.exists(local_driver):
                        service = Service(executable_path=local_driver)
                        driver = webdriver.Chrome(service=service, options=chrome_options)
                        print("Chrome iniciado com driver local")
                    else:
                        raise Exception("ChromeDriver local não encontrado")
                        
                except Exception as e2:
                    print(f"Segunda tentativa falhou: {str(e2)}")
                    try:
                        # Terceira tentativa: procurar o chromedriver no PATH
                        driver = webdriver.Chrome(options=chrome_options)
                        print("Chrome iniciado com driver do PATH")
                    except Exception as e3:
                        print(f"Todas as tentativas falharam:")
                        print(f"1: {str(e1)}")
                        print(f"2: {str(e2)}")
                        print(f"3: {str(e3)}")
                        raise Exception("Não foi possível inicializar o Chrome. Por favor, verifique se o Chrome está instalado e atualizado.")

            if not driver:
                raise Exception("Não foi possível inicializar o Chrome")
            
            wait = WebDriverWait(driver, 10)
            
            url = "https://minha.effecti.com.br/#/proposta-minhas"
            driver.get(url)
            
            try:
                # Login mais rápido e direto
                email_field = wait.until(EC.presence_of_element_located((By.NAME, "input-login")))
                password_field = wait.until(EC.presence_of_element_located((By.NAME, "input-password")))
                
                # Preenche email e senha de uma vez
                email_field.send_keys("fernanda@alcantaramendes.com.br")
                password_field.send_keys("Alcantara@2025")
                
                # Clica no botão Entrar imediatamente
                entrar_button = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "div.button-submit a.login-btn.l-button")))
                entrar_button.click()
                print("Login realizado com sucesso!")
                
                # Reduz o tempo de espera e já procura o botão Cadastrar Proposta
                print("Aguardando página carregar...")
                time.sleep(3)  # Aguarda a página carregar completamente
                
                # Espera o overlay desaparecer
                try:
                    overlay = WebDriverWait(driver, 10).until(
                        EC.invisibility_of_element_located((By.CSS_SELECTOR, "div.overlay.fullscreen"))
                    )
                except:
                    print("Aviso: Overlay não encontrado ou já invisível")
                
                print("Procurando botão Cadastrar Proposta...")
                try:
                    # Primeira tentativa: esperar o botão estar clicável
                    cadastrar_proposta_button = WebDriverWait(driver, 10).until(
                        EC.element_to_be_clickable((By.CSS_SELECTOR, "button.btn-new-proposal"))
                    )
                    cadastrar_proposta_button.click()
                except:
                    try:
                        # Segunda tentativa: usar JavaScript para remover overlay e clicar
                        driver.execute_script("""
                            // Remove qualquer overlay
                            var overlays = document.querySelectorAll('div.overlay');
                            overlays.forEach(function(overlay) {
                                overlay.remove();
                            });
                            
                            // Encontra e clica no botão
                            var botao = document.querySelector('button.btn-new-proposal');
                            if(botao) {
                                botao.click();
                            }
                        """)
                    except:
                        # Terceira tentativa: localizar por XPath e usar Actions
                        from selenium.webdriver.common.action_chains import ActionChains
                        button = driver.find_element(By.XPATH, "//button[contains(@class, 'btn-new-proposal')]")
                        actions = ActionChains(driver)
                        actions.move_to_element(button).click().perform()
                
                print("Iniciando cadastro de proposta...")
                
                # Fecha o popup do tour se aparecer, com tempo reduzido
                if fechar_popup_tour(driver, wait):
                    print("Tour inicial fechado!")
                
                # Carrega os dados do JSON
                with open('dados.json', 'r', encoding='utf-8') as arquivo:
                    dados = json.load(arquivo)
                
                # Seleciona portal e empresa em sequência rápida
                select_element = wait.until(EC.presence_of_element_located((By.ID, "sel-portal")))
                select = Select(select_element)
                
                # Mapeamento dos IDs dos portais
                portal_mapping = {
                    'bll': '24',
                    'comprasnet': '1',
                    'compraspublicas': '3',
                    'bnc': '1362',
                    'comprasbr': '898',
                    'licitanet': '28'
                }
                
                # Seleciona o portal
                portal_id = portal_mapping.get(dados['portal_id'])
                if portal_id:
                    select.select_by_value(portal_id)
                    print(f"Portal selecionado: {dados['portal_nome']}")
                    
                    # Seleciona a empresa Fernanda imediatamente após o portal
                    empresa_select = wait.until(EC.presence_of_element_located((By.ID, "branch")))
                    Select(empresa_select).select_by_value("1")
                    print("Empresa selecionada!")
                    
                    # Mapeamento dos campos específicos por portal
                    field_mapping = {
                        'bll': {
                            'bidding': 'sel-bidding',
                            'organ': 'sel-organ',
                            'quotation': 'sel-quotation'
                        },
                        'comprasnet': {
                            'bidding': 'sel-bidding',
                            'uasg': 'sel-organ',
                            'quotation': 'sel-quotation',
                            'uasg_name': 'sel-uasg-name'
                        },
                        'compraspublicas': {
                            'bidding': 'sel-bidding',
                            'organ': 'sel-organ',
                            'quotation': 'sel-quotation'
                        },
                        'bnc': {
                            'bidding': 'sel-bidding',
                            'organ': 'sel-organ',
                            'quotation': 'sel-quotation'
                        },
                        'comprasbr': {
                            'bidding': 'sel-bidding',
                            'organ': 'sel-organ',
                            'quotation': 'sel-quotation'
                        },
                        'licitanet': {
                            'bidding': 'sel-bidding',
                            'organ': 'sel-organ',
                            'quotation': 'sel-quotation'
                        }
                    }
                    
                    # Pega os IDs corretos para o portal atual
                    portal_fields = field_mapping.get(dados['portal_id'])
                    
                    if portal_fields:
                        fechar_popup_tour(driver, wait)
                        
                        if dados['portal_id'] == 'comprasnet':
                            try:
                                # Campo Licitação (Número da compra)
                                print("Tentando preencher campo Licitação...")
                                bidding_field = wait.until(EC.presence_of_element_located((By.ID, portal_fields['bidding'])))
                                bidding_field.click()
                                
                                # Usar valores padrão se não houver registros
                                if not dados.get('registros'):
                                    print("Nenhum registro encontrado, usando valores padrão...")
                                    edital_numero = '990012024'
                                    uasg = '980425'
                                    modalidade = 'Pregão'
                                else:
                                    edital_numero = dados['registros'][-1].get('numero_compra', '990012024')
                                    uasg = dados['registros'][-1].get('uasg', '980425')
                                    modalidade = dados['registros'][-1].get('modalidade', 'Pregão')
                                
                                bidding_field.send_keys(str(edital_numero))
                                print(f"Campo Licitação preenchido: {edital_numero}")
                                
                                # Campo UASG
                                print("Tentando preencher campo UASG...")
                                uasg_field = wait.until(EC.presence_of_element_located((By.ID, portal_fields['uasg'])))
                                uasg_field.click()
                                uasg_field.send_keys(str(uasg))
                                print(f"Campo UASG preenchido: {uasg}")
                                
                                # Modalidade
                                print("\nTentando selecionar modalidade no portal comprasnet...")
                                select_element = wait.until(EC.presence_of_element_located((By.ID, portal_fields['quotation'])))
                                select = Select(select_element)
                                
                                # Obter todas as opções disponíveis
                                available_options = []
                                for option in select.options:
                                    value = option.get_attribute('value')
                                    text = option.text
                                    available_options.append((value, text))
                                    print(f"Opção encontrada: valor='{value}', texto='{text}'")
                                
                                # Mapear a modalidade para o valor correto
                                modalidade_id = MODALIDADE_VALUES['comprasnet'].get(modalidade)
                                print(f"Modalidade selecionada: {modalidade} (tentando usar valor: {modalidade_id})")
                                
                                if modalidade_id and any(modalidade_id == opt[0] for opt in available_options):
                                    select.select_by_value(modalidade_id)
                                    print(f"Modalidade selecionada com sucesso: {modalidade} (valor: {modalidade_id})")
                                else:
                                    print(f"AVISO: Valor da modalidade '{modalidade}' ({modalidade_id}) não disponível.")
                                    print(f"Valores disponíveis: {available_options}")
                                    print("Usando valor padrão '1'")
                                    select.select_by_value('1')
                            except Exception as e:
                                print(f"Erro ao preencher campos do ComprasNet: {str(e)}")
                                print("Detalhes do erro:")
                                import traceback
                                print(traceback.format_exc())
                                return
                                
                        else:
                            # Preenchimento padrão para outros portais
                            try:
                                # Campo Licitação/Edital
                                bidding_field = wait.until(EC.presence_of_element_located((By.ID, portal_fields['bidding'])))
                                bidding_field.click()
                                edital_numero = dados['registros'][-1].get('numero_edital', '')
                                bidding_field.send_keys(str(edital_numero))
                                print(f"Campo Licitação preenchido: {edital_numero}")
                                
                                # Campo Órgão
                                organ_field = wait.until(EC.presence_of_element_located((By.ID, portal_fields['organ'])))
                                organ_field.click()
                                orgao = dados['registros'][-1].get('orgao', '')
                                organ_field.send_keys(str(orgao))
                                print(f"Campo Órgão preenchido: {orgao}")
                                
                                # Modalidade
                                print("\nTentando selecionar modalidade no portal...")
                                select_element = wait.until(EC.presence_of_element_located((By.ID, portal_fields['quotation'])))
                                select = Select(select_element)
                                
                                # Obter todas as opções disponíveis
                                available_options = []
                                for option in select.options:
                                    value = option.get_attribute('value')
                                    text = option.text
                                    available_options.append((value, text))
                                    print(f"Opção encontrada: valor='{value}', texto='{text}'")
                                
                                # Mapear a modalidade para o valor correto
                                modalidade = dados['registros'][-1].get('modalidade', 'Pregão')
                                modalidade_id = MODALIDADE_VALUES[dados['portal_id']].get(modalidade)
                                print(f"Modalidade selecionada: {modalidade} (tentando usar valor: {modalidade_id})")
                                
                                if modalidade_id and any(modalidade_id == opt[0] for opt in available_options):
                                    select.select_by_value(modalidade_id)
                                    print(f"Modalidade selecionada com sucesso: {modalidade} (valor: {modalidade_id})")
                                else:
                                    print(f"AVISO: Valor da modalidade '{modalidade}' ({modalidade_id}) não disponível.")
                                    print(f"Valores disponíveis: {available_options}")
                                    print("Usando valor padrão '1'")
                                    select.select_by_value('1')
                                
                            except Exception as e:
                                print(f"Erro ao preencher campos: {str(e)}")
                                return
                        
                        # Carregar Itens
                        print("Clicando em Carregar Itens...")
                        carregar_itens_button = wait.until(EC.element_to_be_clickable((By.XPATH, "//button[contains(text(), 'Carregar Itens')]")))
                        carregar_itens_button.click()
                        print("Botão Carregar Itens clicado!")
                        
                        time.sleep(2)  # Aguarda carregamento
                        
                        # Verificar aviso
                        if verificar_aviso(driver, wait):
                            print("Aviso detectado e fechado. Continuando com a exportação...")
                            time.sleep(1)  # Pequena pausa após fechar o aviso
                        
                        # Scroll e exportação
                        print("Rolando a página...")
                        driver.execute_script("window.scrollBy(0, 300);")
                        time.sleep(1)  # Aguarda o scroll

                        # Após rolar, tenta fechar o span do tour
                        try:
                            print("Procurando botão finalizar do tour...")
                            finalizar_button = driver.find_element(By.CSS_SELECTOR, "div.popover.minha-tour.tour-tour-proposta-cadastro a[data-role='end']")
                            finalizar_button.click()
                            print("Tour fechado com sucesso!")
                            time.sleep(1)  # Aguarda o tour fechar
                        except Exception as e:
                            print(f"Aviso: Tour não encontrado após scroll - {str(e)}")
                        
                        # Exportar planilha
                        try:
                            print("Procurando botão Exportar planilha...")
                            exportar = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 'button[title="Exportar planilha"]')))
                            driver.execute_script("arguments[0].click();", exportar)
                            print("Botão Exportar planilha clicado!")
                            
                            time.sleep(2)  # Aguarda modal abrir
                            
                            print("Procurando botão Exportar na modal...")
                            # Tenta diferentes seletores para o botão Exportar
                            try:
                                # Primeira tentativa - botão pela classe
                                confirmar_exportar = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "div.modal-footer button.btn-primary")))
                            except:
                                try:
                                    # Segunda tentativa - botão pelo texto
                                    confirmar_exportar = wait.until(EC.element_to_be_clickable((By.XPATH, "//button[contains(text(), 'Exportar')]")))
                                except:
                                    # Terceira tentativa - botão pela estrutura HTML fornecida
                                    confirmar_exportar = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "div[data-v-1d2d804c] button.btn.btn-primary")))
                            
                            # Tenta clicar de diferentes formas
                            try:
                                confirmar_exportar.click()
                            except:
                                try:
                                    driver.execute_script("arguments[0].click();", confirmar_exportar)
                                except:
                                    # Última tentativa - força o clique via JavaScript
                                    driver.execute_script("""
                                        var buttons = document.querySelectorAll('button');
                                        for(var i = 0; i < buttons.length; i++) {
                                            if(buttons[i].textContent.includes('Exportar')) {
                                                buttons[i].click();
                                                break;
                                            }
                                        }
                                    """)
                            
                            print("Exportação iniciada!")
                            
                            # Aguarda o download do arquivo
                            print("Aguardando download do arquivo...")
                            download_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'downloads')
                            timeout = 30
                            start_time = time.time()
                            
                            while time.time() - start_time < timeout:
                                # Procura por arquivos que começam com 'exportacao_' e terminam com '.xlsx'
                                files = [f for f in os.listdir(download_dir) if f.startswith('exportação_') and f.endswith('.xlsx')]
                                if files:
                                    # Pega o arquivo mais recente
                                    newest_file = max([os.path.join(download_dir, f) for f in files], key=os.path.getctime)
                                    print(f"Arquivo exportado encontrado: {newest_file}")
                                    time.sleep(2)  # Garante que o arquivo terminou de ser baixado)
                                    
                                    # Após encontrar o arquivo, inicia o processamento
                                    print("\nIniciando processamento do arquivo...")
                                    from excel_processor import processar_arquivo_exportado
                                    
                                    resultado = self.processar_arquivo_exportado(newest_file)
                                    if resultado and isinstance(resultado, dict):
                                        print("\nArquivos gerados:")
                                        print(f"Excel processado: {resultado.get('excel', 'Não gerado')}")
                                        print(f"JSON gerado: {resultado.get('json', 'Não gerado')}")
                                        
                                        # Remove o arquivo Excel original após processamento
                                        try:
                                            os.remove(newest_file)
                                            print(f"\nArquivo Excel original removido: {newest_file}")
                                        except Exception as e:
                                            print(f"Aviso: Não foi possível remover o arquivo original: {str(e)}")
                                        
                                        # Gera o documento Word
                                        from word_processor import criar_tabela_word
                                        template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'word', 'MUNDIAL PROPOSTA 447338.docx')
                                        output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'proposta_final.docx')
                                        
                                        # Tenta com diferentes variações do marcador
                                        marcadores = [
                                            "{{TABELA_AQUI}}",
                                            "{{tabela_aqui}}",
                                            "{TABELA_AQUI}",
                                            "{tabela_aqui}",
                                            "TABELA_AQUI",
                                            "tabela_aqui"
                                        ]
                                        
                                        documento_word = None
                                        for marcador in marcadores:
                                            print(f"\nTentando com marcador: {marcador}")
                                            documento_word = criar_tabela_word(
                                                json_file=resultado['json'],
                                                template_file=template_path,
                                                marcador_tabela=marcador,
                                                output_file=output_path
                                            )
                                            if documento_word:
                                                print(f"Documento Word gerado com sucesso usando marcador: {marcador}")
                                                break
                                        
                                        if documento_word:
                                            print(f"Documento Word gerado: {documento_word}")
                                        else:
                                            print("Não foi possível gerar o documento Word com nenhum dos marcadores tentados")
                                        
                                        print("\nProcessamento concluído com sucesso!")
                                        return {
                                            'excel': resultado['excel'],
                                            'json': resultado['json'],
                                            'word': documento_word
                                        }
                                    else:
                                        print("Erro: Não foi possível processar o arquivo Excel")
                                    
                                time.sleep(1)
                            
                            print("Tempo limite de download excedido!")
                            
                        except Exception as e:
                            print(f"Erro durante a exportação: {str(e)}")
                            print("Detalhes do erro:")
                            import traceback
                            print(traceback.format_exc())
                    
            except Exception as e:
                print(f"Erro durante o processo: {str(e)}")
                print("Detalhes do erro:")
                import traceback
                print(traceback.format_exc())
                
        except Exception as e:
            print(f"Erro ao inicializar o Chrome: {str(e)}")
            print("Detalhes do erro:")
            import traceback
            print(traceback.format_exc())
            
        finally:
            if driver:
                driver.quit()
    
    def encontrar_ultimo_excel(self):
        """Encontra o arquivo Excel mais recente na pasta de downloads."""
        downloads_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'downloads')
        # Procura por arquivos que começam com 'exportacao_' e terminam com .xlsx
        lista_arquivos = glob.glob(os.path.join(downloads_path, 'exportacao_*.xlsx'))
        if not lista_arquivos:
            return None
        # Retorna o arquivo mais recente
        return max(lista_arquivos, key=os.path.getctime)

    def processar_excel(self, arquivo_excel):
        """Processa o arquivo Excel, modificando valores específicos."""
        try:
            # Lê o arquivo original com openpyxl engine
            print("Iniciando leitura do arquivo Excel...")
            df = pd.read_excel(arquivo_excel, engine='openpyxl')
            new_df = df.copy()

            print(f"Arquivo lido com sucesso. Colunas encontradas: {list(df.columns)}")

            # Modifica valores da terceira coluna
            if len(df.columns) > 2:
                third_column = df.columns[2]
                new_df.iloc[1:, 2] = 0
                print(f"Valores da terceira coluna ({third_column}) modificados")

            # Preenche campos "marca" e "modelo"
            colunas_lower = [col.lower() for col in new_df.columns]
            if 'marca' in colunas_lower:
                idx = colunas_lower.index('marca')
                new_df.iloc[:, idx] = 'Própria'
                print("Campo 'marca' preenchido")
            if 'modelo' in colunas_lower:
                idx = colunas_lower.index('modelo')
                new_df.iloc[:, idx] = 'Própria'
                print("Campo 'modelo' preenchido")

            # Cria o nome do arquivo modificado
            nome_base = os.path.splitext(os.path.basename(arquivo_excel))[0]
            novo_arquivo = arquivo_excel.replace('.xlsx', '_modificado.xlsx')
            
            # Salva o arquivo modificado
            print(f"Salvando arquivo modificado como: {novo_arquivo}")
            new_df.to_excel(novo_arquivo, index=False, engine='openpyxl')
            print(f"Excel processado e salvo com sucesso")
            return novo_arquivo
        except Exception as e:
            print(f"Erro ao processar Excel: {str(e)}")
            print("Detalhes do erro:")
            import traceback
            print(traceback.format_exc())
            return None

    def excel_para_json(self, arquivo_excel):
        """Converte o arquivo Excel para JSON com formato específico."""
        try:
            # Lê o arquivo Excel com openpyxl engine
            print("Iniciando conversão do Excel para JSON...")
            df = pd.read_excel(arquivo_excel, engine='openpyxl')
            
            # Prepara os dados no formato desejado
            nome_base = os.path.splitext(os.path.basename(arquivo_excel))[0]
            data = {
                "processo": nome_base,
                "colunas": list(df.columns),
                "dados": df.to_dict(orient="records")
            }
            
            # Cria o nome do arquivo JSON
            arquivo_json = arquivo_excel.replace('.xlsx', '.json')
            
            # Salva o JSON
            print(f"Salvando dados em: {arquivo_json}")
            with open(arquivo_json, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=4)
            
            print(f"Arquivo JSON criado com sucesso")
            return arquivo_json
        except Exception as e:
            print(f"Erro ao converter Excel para JSON: {str(e)}")
            print("Detalhes do erro:")
            import traceback
            print(traceback.format_exc())
            return None

    def excluir_excel_temporario(self, arquivo_excel):
        """Exclui o arquivo Excel após a conversão para JSON."""
        try:
            if os.path.exists(arquivo_excel):
                os.remove(arquivo_excel)
                print(f"Arquivo Excel excluído: {arquivo_excel}")
        except Exception as e:
            print(f"Erro ao excluir arquivo Excel: {str(e)}")

    def processar_arquivo_exportado(self, arquivo_excel):
        """
        Processa o arquivo Excel exportado e prepara os dados para o Word
        """
        try:
            print(f"\nProcessando arquivo Excel: {arquivo_excel}")
            # Lê o arquivo Excel
            df = pd.read_excel(arquivo_excel, engine='openpyxl')
            
            # Remove linhas vazias
            df = df.dropna(how='all')
            
            # Processa as colunas necessárias
            dados_processados = []
            for index, row in df.iterrows():
                item = {
                    'item': row.get('Item', ''),
                    'descricao': row.get('Descrição', ''),
                    'unidade': row.get('Unidade de fornecimento', ''),
                    'quantidade': row.get('Quantidade', 0),
                    'valor_unitario': row.get('Valor Unitário', 0),
                    'valor_total': row.get('Valor Total', 0)
                }
                dados_processados.append(item)
            
            print(f"Processados {len(dados_processados)} itens do Excel")
            return dados_processados
            
        except Exception as e:
            print(f"Erro ao processar arquivo Excel: {str(e)}")
            print("Detalhes do erro:")
            import traceback
            print(traceback.format_exc())
            return None

    def gerar_documento_word(self, dados_processados, template_word=None):
        """
        Gera um documento Word com os dados processados
        """
        try:
            from docx import Document
            from docx.shared import Pt
            
            # Se não fornecido template, cria novo documento
            if template_word and os.path.exists(template_word):
                doc = Document(template_word)
                print(f"Usando template: {template_word}")
            else:
                doc = Document()
                print("Criando novo documento Word")
            
            # Adiciona tabela com os dados
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            
            # Cabeçalho
            header_cells = table.rows[0].cells
            headers = ['Item', 'Descrição', 'Unidade', 'Quantidade', 'Valor Unitário', 'Valor Total']
            for i, header in enumerate(headers):
                header_cells[i].text = header
            
            # Dados
            for item in dados_processados:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item['item'])
                row_cells[1].text = str(item['descricao'])
                row_cells[2].text = str(item['unidade'])
                row_cells[3].text = str(item['quantidade'])
                row_cells[4].text = f"R$ {item['valor_unitario']:.2f}"
                row_cells[5].text = f"R$ {item['valor_total']:.2f}"
            
            # Salva o documento
            output_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'proposta_processada.docx')
            doc.save(output_file)
            print(f"\nDocumento Word gerado com sucesso: {output_file}")
            return output_file
            
        except Exception as e:
            print(f"Erro ao gerar documento Word: {str(e)}")
            print("Detalhes do erro:")
            import traceback
            print(traceback.format_exc())
            return None

    def processar_arquivo(self, arquivo_excel, arquivo_word):
        """
        Processa o arquivo Excel e gera o documento Word
        """
        try:
            print("\nIniciando processamento do arquivo...")
            
            # Processa o Excel e gera o JSON
            resultado = self.processar_arquivo_exportado(arquivo_excel)
            if not resultado:
                print("Erro ao processar o arquivo Excel")
                return False
            
            # Remove o arquivo Excel processado pois não será mais necessário
            try:
                os.remove(resultado['excel'])
                print(f"\nArquivo Excel original removido: {resultado['excel']}")
            except:
                pass
            
            # Lista de marcadores para tentar
            marcadores = [
                "{{TABELA_AQUI}}",
                "{{tabela_aqui}}",
                "{TABELA_AQUI}",
                "{tabela_aqui}",
                "TABELA_AQUI",
                "tabela_aqui"
            ]
            
            # Tenta gerar o documento Word com cada marcador
            for marcador in marcadores:
                print(f"\nTentando com marcador: {marcador}")
                
                # Define o nome do arquivo de saída baseado no template
                template_dir = os.path.dirname(arquivo_word)
                template_nome = os.path.basename(arquivo_word)
                nome_base = os.path.splitext(template_nome)[0]
                arquivo_saida = os.path.join(template_dir, f"{nome_base}_proposta.docx")
                
                # Gera o documento Word
                doc_gerado = criar_tabela_word(
                    resultado['json'],
                    template_file=arquivo_word,
                    marcador_tabela=marcador,
                    output_file=arquivo_saida
                )
                
                if doc_gerado:
                    print(f"Documento Word gerado com sucesso usando marcador: {marcador}")
                    print(f"Documento Word gerado: {doc_gerado}")
                    return True
                
            print("\nNão foi possível gerar o documento Word com nenhum dos marcadores tentados")
            return False
        
        except Exception as e:
            print(f"Erro durante o processamento: {str(e)}")
            return False
        finally:
            print("\nProcessamento concluído!")

    def voltar(self):
        self.root.destroy()
        SeletorPortal()

class FormularioBLL(FormularioBase):
    def __init__(self, portal_id, portal_nome):
        super().__init__(portal_id, portal_nome)
        
        self.numero_edital = self.criar_campo(self.campos_frame, "N° do Edital:")
        self.orgao = self.criar_campo(self.campos_frame, "Órgão:")
        self.modalidade = self.criar_campo(self.campos_frame, "Modalidade:", "combobox")
        
        # Configurar as opções da modalidade
        if isinstance(self.modalidade, ttk.Combobox):
            self.modalidade['values'] = MODALIDADES.get(self.portal_id, [])
            self.modalidade.set('Selecione a modalidade')  # Valor padrão
        
        self.root.mainloop()
    
    def formato_lista(self, registro):
        return f"Edital: {registro['numero_edital']} - Órgão: {registro['orgao']} - Modalidade: {registro['modalidade']}"
    
    def coletar_dados(self):
        return {
            "numero_edital": self.numero_edital.get(),
            "orgao": self.orgao.get(),
            "modalidade": self.modalidade.get(),
            "data_cadastro": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
    
    def limpar_campos(self):
        self.numero_edital.delete(0, tk.END)
        self.orgao.delete(0, tk.END)
        self.modalidade.set('Selecione a modalidade')

class FormularioComprasNet(FormularioBase):
    def __init__(self, portal_id, portal_nome):
        super().__init__(portal_id, portal_nome)
        
        self.numero_compra = self.criar_campo(self.campos_frame, "N° da Compra:")
        self.uasg = self.criar_campo(self.campos_frame, "UASG:")
        self.modalidade = self.criar_campo(self.campos_frame, "Modalidade:", "combobox")
        
        self.root.mainloop()
    
    def formato_lista(self, registro):
        return f"Compra: {registro['numero_compra']} - UASG: {registro['uasg']} - Modalidade: {registro['modalidade']}"
    
    def coletar_dados(self):
        return {
            "numero_compra": self.numero_compra.get(),
            "uasg": self.uasg.get(),
            "modalidade": self.modalidade.get(),
            "data_cadastro": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
    
    def limpar_campos(self):
        self.numero_compra.delete(0, tk.END)
        self.uasg.delete(0, tk.END)
        self.modalidade.set('Selecione a modalidade')

class FormularioComprasPublicas(FormularioBase):
    def __init__(self, portal_id, portal_nome):
        super().__init__(portal_id, portal_nome)
        
        self.numero_processo = self.criar_campo(self.campos_frame, "N° do Processo:")
        self.orgao = self.criar_campo(self.campos_frame, "Órgão:")
        self.modalidade = self.criar_campo(self.campos_frame, "Modalidade:", "combobox")
        
        self.root.mainloop()
    
    def formato_lista(self, registro):
        return f"Processo: {registro['numero_processo']} - Órgão: {registro['orgao']} - Modalidade: {registro['modalidade']}"
    
    def coletar_dados(self):
        return {
            "numero_processo": self.numero_processo.get(),
            "orgao": self.orgao.get(),
            "modalidade": self.modalidade.get(),
            "data_cadastro": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
    
    def limpar_campos(self):
        self.numero_processo.delete(0, tk.END)
        self.orgao.delete(0, tk.END)
        self.modalidade.set('Selecione a modalidade')

class FormularioBNC(FormularioBase):
    def __init__(self, portal_id, portal_nome):
        super().__init__(portal_id, portal_nome)
        
        self.numero_edital = self.criar_campo(self.campos_frame, "N° do Edital:")
        self.orgao = self.criar_campo(self.campos_frame, "Órgão:")
        self.modalidade = self.criar_campo(self.campos_frame, "Modalidade:", "combobox")
        
        # Configurar as opções da modalidade
        if isinstance(self.modalidade, ttk.Combobox):
            self.modalidade['values'] = MODALIDADES.get(self.portal_id, [])
            self.modalidade.set('Selecione a modalidade')  # Valor padrão
        
        self.root.mainloop()
    
    def formato_lista(self, registro):
        return f"Edital: {registro['numero_edital']} - Órgão: {registro['orgao']} - Modalidade: {registro['modalidade']}"
    
    def coletar_dados(self):
        return {
            "numero_edital": self.numero_edital.get(),
            "orgao": self.orgao.get(),
            "modalidade": self.modalidade.get(),
            "data_cadastro": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
    
    def limpar_campos(self):
        self.numero_edital.delete(0, tk.END)
        self.orgao.delete(0, tk.END)
        self.modalidade.set('Selecione a modalidade')

class FormularioComprasBR(FormularioBase):
    def __init__(self, portal_id, portal_nome):
        super().__init__(portal_id, portal_nome)
        
        self.numero_edital = self.criar_campo(self.campos_frame, "N° do Edital:")
        self.orgao = self.criar_campo(self.campos_frame, "Órgão:")
        self.modalidade = self.criar_campo(self.campos_frame, "Modalidade:", "combobox")
        
        self.root.mainloop()
    
    def formato_lista(self, registro):
        return f"Edital: {registro['numero_edital']} - Órgão: {registro['orgao']} - Modalidade: {registro['modalidade']}"
    
    def coletar_dados(self):
        return {
            "numero_edital": self.numero_edital.get(),
            "orgao": self.orgao.get(),
            "modalidade": self.modalidade.get(),
            "data_cadastro": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
    
    def limpar_campos(self):
        self.numero_edital.delete(0, tk.END)
        self.orgao.delete(0, tk.END)
        self.modalidade.set('Selecione a modalidade')

class FormularioLicitaNet(FormularioBase):
    def __init__(self, portal_id, portal_nome):
        super().__init__(portal_id, portal_nome)
        
        self.numero_edital = self.criar_campo(self.campos_frame, "N° do Edital:")
        self.orgao = self.criar_campo(self.campos_frame, "Órgão:")
        self.modalidade = self.criar_campo(self.campos_frame, "Modalidade:", "combobox")
        
        self.root.mainloop()
    
    def formato_lista(self, registro):
        return f"Edital: {registro['numero_edital']} - Órgão: {registro['orgao']} - Modalidade: {registro['modalidade']}"
    
    def coletar_dados(self):
        return {
            "numero_edital": self.numero_edital.get(),
            "orgao": self.orgao.get(),
            "modalidade": self.modalidade.get(),
            "data_cadastro": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
    
    def limpar_campos(self):
        self.numero_edital.delete(0, tk.END)
        self.orgao.delete(0, tk.END)
        self.modalidade.set('Selecione a modalidade')

class TelaDetalhes:
    def __init__(self, dados_processo):
        self.root = tk.Toplevel()
        self.root.title("Detalhes do Processo")
        self.root.geometry("1000x600")
        
        # Frame principal
        main_frame = tk.Frame(self.root)
        main_frame.pack(expand=True, fill='both', padx=20, pady=20)
        
        # Título
        titulo = tk.Label(
            main_frame,
            text="Detalhes Completos do Processo",
            font=('Helvetica', 18, 'bold')
        )
        titulo.pack(pady=20)
        
        # Frame para a árvore de dados
        tree_frame = tk.Frame(main_frame)
        tree_frame.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Scrollbars
        vsb = ttk.Scrollbar(tree_frame, orient="vertical")
        hsb = ttk.Scrollbar(tree_frame, orient="horizontal")
        
        # Treeview para mostrar os dados em formato de árvore
        self.tree = ttk.Treeview(tree_frame, selectmode='browse',
                                yscrollcommand=vsb.set,
                                xscrollcommand=hsb.set)
        
        # Configurar scrollbars
        vsb.config(command=self.tree.yview)
        hsb.config(command=self.tree.xview)
        
        # Posicionar elementos
        vsb.pack(side='right', fill='y')
        hsb.pack(side='bottom', fill='x')
        self.tree.pack(fill='both', expand=True)
        
        # Configurar colunas
        self.tree["columns"] = ("Valor",)
        self.tree.column("#0", width=300, stretch=tk.NO)
        self.tree.column("Valor", width=500, stretch=tk.NO)
        
        self.tree.heading("#0", text="Campo")
        self.tree.heading("Valor", text="Valor")
        
        # Preencher a árvore com os dados
        self.preencher_arvore(dados_processo)
    
    def preencher_arvore(self, dados, parent=""):
        if isinstance(dados, dict):
            for key, value in dados.items():
                item_id = self.tree.insert(parent, "end", text=str(key))
                if isinstance(value, (dict, list)):
                    self.preencher_arvore(value, item_id)
                else:
                    self.tree.insert(parent, "end", text=str(key), values=(str(value),))
        elif isinstance(dados, list):
            for i, item in enumerate(dados):
                item_id = self.tree.insert(parent, "end", text=f"Item {i+1}")
                if isinstance(item, (dict, list)):
                    self.preencher_arvore(item, item_id)
                else:
                    self.tree.insert(parent, "end", text=str(item))

class TelaFila:
    def __init__(self, dados_json):
        self.root = tk.Toplevel()
        self.root.title("Fila de Processos")
        self.root.geometry("800x600")
        
        # Frame principal
        main_frame = tk.Frame(self.root)
        main_frame.pack(expand=True, fill='both', padx=20, pady=20)
        
        # Título
        titulo = tk.Label(
            main_frame,
            text="Fila de Processos Aceitos",
            font=('Helvetica', 18, 'bold')
        )
        titulo.pack(pady=20)
        
        # Frame para lista de cards
        self.cards_frame = tk.Frame(main_frame)
        self.cards_frame.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Lista de cards com scrollbar
        cards_scrollbar = tk.Scrollbar(self.cards_frame)
        cards_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.cards_listbox = tk.Listbox(
            self.cards_frame,
            height=10,
            width=50,
            yscrollcommand=cards_scrollbar.set,
            font=('Helvetica', 10)
        )
        self.cards_listbox.pack(fill='both', expand=True)
        cards_scrollbar.config(command=self.cards_listbox.yview)
        
        # Botão para ver detalhes
        self.btn_detalhes = tk.Button(
            main_frame,
            text="Ver Detalhes",
            command=self.ver_detalhes,
            font=('Helvetica', 12),
            width=20,
            height=2
        )
        self.btn_detalhes.pack(pady=20)
        
        # Armazena os dados JSON
        self.dados_json = dados_json
        self.cards_data = []  # Armazena os dados completos dos cards
        
        # Carregar cards aceitos
        self.carregar_cards_aceitos()
        
        # Bind do duplo clique
        self.cards_listbox.bind('<Double-Button-1>', lambda e: self.ver_detalhes())
    
    def carregar_cards_aceitos(self):
        try:
            self.cards_data = []  # Limpa os dados anteriores
            
            for empresa_id, info in self.dados_json.items():
                if "etapas" in info and "ACEITAS" in info["etapas"]:
                    nome_empresa = info["nome_empresa"]
                    for card in info["etapas"]["ACEITAS"]["cards"]:
                        numero_pregao = card.get("Número do pregão", "N/A")
                        portal = card.get("Portal", "N/A")
                        # Armazena o card completo e o texto para display
                        self.cards_data.append({
                            'texto': f"{nome_empresa} - Pregão: {numero_pregao} - Portal: {portal}",
                            'dados': {
                                'empresa_id': empresa_id,
                                'empresa': info,
                                'card': card
                            }
                        })
            
            # Limpa a lista atual
            self.cards_listbox.delete(0, tk.END)
            
            # Adiciona os cards encontrados
            for card in self.cards_data:
                self.cards_listbox.insert(tk.END, card['texto'])
            
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao carregar cards: {str(e)}")
    
    def ver_detalhes(self):
        selection = self.cards_listbox.curselection()
        if selection:
            index = selection[0]
            dados_card = self.cards_data[index]['dados']
            TelaDetalhes(dados_card)

class TelaPrincipal:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Sistema de Propostas")
        self.root.geometry("1600x900")
        
        # Definir cores
        self.bg_color = "#f0f0f0"
        self.frame_bg = "#ffffff"
        self.highlight_color = "#e1e1e1"
        
        self.root.configure(bg=self.bg_color)
        
        # Frame principal que conterá os três painéis
        self.main_frame = tk.Frame(self.root, bg=self.bg_color)
        self.main_frame.pack(expand=True, fill='both', padx=10, pady=10)
        
        # Configurar o grid com pesos iguais
        self.main_frame.grid_columnconfigure(0, weight=1)  # Controle: 33%
        self.main_frame.grid_columnconfigure(1, weight=1)  # Fila: 33%
        self.main_frame.grid_columnconfigure(2, weight=1)  # Detalhes: 33%
        self.main_frame.grid_rowconfigure(0, weight=1)
        
        # Criar os três painéis
        self.criar_painel_controle()  # Painel esquerdo
        self.criar_painel_fila()      # Painel central
        self.criar_painel_detalhes()  # Painel direito
        
        # Inicializar dados
        self.dados_json = None
        self.cards_data = []
        self.output_buffer = []
        self.processando = False
        self.parar_processamento = False
        
    def criar_painel_controle(self):
        # Painel de Controle (Esquerdo)
        controle_frame = tk.LabelFrame(
            self.main_frame,
            text="Painel de Controle",
            padx=10, pady=10,
            bg=self.frame_bg,
            font=('Helvetica', 10, 'bold')
        )
        controle_frame.grid(row=0, column=0, sticky='nsew', padx=5)
        
        # Frame para os botões
        btn_frame = tk.Frame(controle_frame, bg=self.frame_bg)
        btn_frame.pack(fill='x', pady=10)
        
        # Botão Iniciar
        self.btn_iniciar = tk.Button(
            btn_frame,
            text="Iniciar Processamento",
            command=self.iniciar_processamento,
            bg='#4CAF50',
            fg='white',
            font=('Helvetica', 10, 'bold'),
            relief=tk.RAISED,
            width=20,
            height=2
        )
        self.btn_iniciar.pack(side='left', padx=5)
        
        # Botão Parar
        self.btn_parar = tk.Button(
            btn_frame,
            text="Parar Processamento",
            command=self.parar_processamento_handler,
            bg='#f44336',
            fg='white',
            font=('Helvetica', 10, 'bold'),
            relief=tk.RAISED,
            width=20,
            height=2,
            state=tk.DISABLED
        )
        self.btn_parar.pack(side='left', padx=5)
        
        # Status
        self.status_label = tk.Label(
            controle_frame,
            text="Status: Aguardando",
            font=('Helvetica', 10),
            bg=self.frame_bg
        )
        self.status_label.pack(pady=10)
        
        # Frame para o log
        log_frame = tk.Frame(controle_frame, bg=self.frame_bg)
        log_frame.pack(fill='both', expand=True)
        
        # Área de log
        self.log_area = tk.Text(
            log_frame,
            height=10,
            font=('Consolas', 10),
            wrap=tk.WORD
        )
        self.log_area.pack(fill='both', expand=True)
        
    def parar_processamento_handler(self):
        self.parar_processamento = True
        self.btn_parar.config(state=tk.DISABLED)
        self.adicionar_log("Solicitação de parada recebida. Aguardando finalização do processo atual...")
        self.adicionar_output("Parando processamento...", "error")
        
    def processar_proximo_item(self):
        if self.parar_processamento:
            self.finalizar_processamento("Processamento interrompido pelo usuário")
            return
            
        if not self.cards_data:
            self.finalizar_processamento("Fila vazia - processamento concluído")
            return
            
        # Pegar o primeiro item da fila
        item = self.cards_data[0]
        self.cards_data = self.cards_data[1:]  # Remove o primeiro item
        
        # Atualizar lista
        self.atualizar_lista_cards()
        self.atualizar_contador()
        
        # Debug dos dados do item
        self.adicionar_output("\nDados do item:", "info")
        for key, value in item.items():
            self.adicionar_output(f"{key}: {value}", "info")
        
        # Identificar portal
        portal = item.get('portal', '').lower()
        self.adicionar_log(f"Processando item do portal: {portal}")
        self.adicionar_output(f"\nIniciando processamento do item: {item.get('id', 'N/A')}", "info")
        self.adicionar_output(f"Portal: {portal}", "info")
        
        # Processar baseado no portal
        try:
            if 'portal1' in portal:
                self.processar_portal1(item)
            elif 'portal2' in portal:
                self.processar_portal2(item)
            else:
                self.adicionar_output(f"Portal não reconhecido: {portal}", "error")
            
            # Agendar próximo item
            self.root.after(1000, self.processar_proximo_item)
            
        except Exception as e:
            self.adicionar_output(f"Erro ao processar item: {str(e)}", "error")
            # Mesmo com erro, continua para o próximo
            self.root.after(1000, self.processar_proximo_item)
    
    def processar_portal1(self, item):
        """Processa um item do Portal 1"""
        try:
            self.adicionar_log(f"Processando Portal 1 - Pregão: {item['pregao']}")
            
            # Instanciar o navegador
            driver = self.iniciar_chrome()
            if not driver:
                raise Exception("Erro ao iniciar o Chrome")
            
            try:
                # Navegar para o portal
                driver.get("URL_DO_PORTAL1")
                
                # Fazer login se necessário
                self.fazer_login_portal1(driver)
                
                # Buscar o pregão
                self.buscar_pregao_portal1(driver, item['pregao'])
                
                # Preencher proposta
                self.preencher_proposta_portal1(driver, item)
                
                # Enviar proposta
                self.enviar_proposta_portal1(driver)
                
                return True
                
            finally:
                # Sempre fechar o navegador ao finalizar
                driver.quit()
                
        except Exception as e:
            self.adicionar_log(f"Erro ao processar Portal 1: {str(e)}")
            return False
            
    def processar_portal2(self, item):
        """Processa um item do Portal 2"""
        try:
            self.adicionar_log(f"Processando Portal 2 - Pregão: {item['pregao']}")
            
            # Instanciar o navegador
            driver = self.iniciar_chrome()
            if not driver:
                raise Exception("Erro ao iniciar o Chrome")
            
            try:
                # Navegar para o portal
                driver.get("URL_DO_PORTAL2")
                
                # Fazer login se necessário
                self.fazer_login_portal2(driver)
                
                # Buscar o pregão
                self.buscar_pregao_portal2(driver, item['pregao'])
                
                # Preencher proposta
                self.preencher_proposta_portal2(driver, item)
                
                # Enviar proposta
                self.enviar_proposta_portal2(driver)
                
                return True
                
            finally:
                # Sempre fechar o navegador ao finalizar
                driver.quit()
                
        except Exception as e:
            self.adicionar_log(f"Erro ao processar Portal 2: {str(e)}")
            return False
            
    def iniciar_chrome(self):
        """Inicia uma instância do Chrome com as configurações necessárias"""
        try:
            chrome_options = webdriver.ChromeOptions()
            
            # Adicionar argumentos necessários
            chrome_options.add_argument('--start-maximized')
            chrome_options.add_argument('--disable-notifications')
            
            # Iniciar o driver
            driver = webdriver.Chrome(options=chrome_options)
            driver.implicitly_wait(10)  # Espera implícita de 10 segundos
            
            return driver
            
        except Exception as e:
            self.adicionar_log(f"Erro ao iniciar Chrome: {str(e)}")
            return None
            
    def fazer_login_portal1(self, driver):
        """Faz login no Portal 1"""
        try:
            # Encontrar campos de login
            usuario = driver.find_element(By.ID, "username")
            senha = driver.find_element(By.ID, "password")
            
            # Preencher credenciais
            usuario.send_keys("SEU_USUARIO")
            senha.send_keys("SUA_SENHA")
            
            # Clicar no botão de login
            botao_login = driver.find_element(By.ID, "login-button")
            botao_login.click()
            
            # Esperar login completar
            WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((By.CLASS_NAME, "logged-in"))
            )
            
        except Exception as e:
            raise Exception(f"Erro no login Portal 1: {str(e)}")
            
    def fazer_login_portal2(self, driver):
        """Faz login no Portal 2"""
        try:
            # Encontrar campos de login
            usuario = driver.find_element(By.ID, "user")
            senha = driver.find_element(By.ID, "pass")
            
            # Preencher credenciais
            usuario.send_keys("SEU_USUARIO")
            senha.send_keys("SUA_SENHA")
            
            # Clicar no botão de login
            botao_login = driver.find_element(By.ID, "btnLogin")
            botao_login.click()
            
            # Esperar login completar
            WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((By.CLASS_NAME, "user-logged"))
            )
            
        except Exception as e:
            raise Exception(f"Erro no login Portal 2: {str(e)}")
            
    def buscar_pregao_portal1(self, driver, pregao):
        """Busca um pregão específico no Portal 1"""
        try:
            # Encontrar campo de busca
            busca = driver.find_element(By.ID, "search-pregao")
            busca.clear()
            busca.send_keys(pregao)
            
            # Clicar no botão de busca
            botao_busca = driver.find_element(By.ID, "btn-search")
            botao_busca.click()
            
            # Esperar resultado aparecer
            WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((By.CLASS_NAME, "pregao-result"))
            )
            
        except Exception as e:
            raise Exception(f"Erro ao buscar pregão no Portal 1: {str(e)}")
            
    def buscar_pregao_portal2(self, driver, pregao):
        """Busca um pregão específico no Portal 2"""
        try:
            # Encontrar campo de busca
            busca = driver.find_element(By.ID, "pregao-search")
            busca.clear()
            busca.send_keys(pregao)
            
            # Clicar no botão de busca
            botao_busca = driver.find_element(By.ID, "search-button")
            botao_busca.click()
            
            # Esperar resultado aparecer
            WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((By.CLASS_NAME, "result-item"))
            )
            
        except Exception as e:
            raise Exception(f"Erro ao buscar pregão no Portal 2: {str(e)}")
            
    def preencher_proposta_portal1(self, driver, item):
        """Preenche a proposta no Portal 1"""
        try:
            # Preencher os campos necessários
            # (Implementar conforme os campos específicos do portal)
            pass
            
        except Exception as e:
            raise Exception(f"Erro ao preencher proposta no Portal 1: {str(e)}")
            
    def preencher_proposta_portal2(self, driver, item):
        """Preenche a proposta no Portal 2"""
        try:
            # Preencher os campos necessários
            # (Implementar conforme os campos específicos do portal)
            pass
            
        except Exception as e:
            raise Exception(f"Erro ao preencher proposta no Portal 2: {str(e)}")
            
    def enviar_proposta_portal1(self, driver):
        """Envia a proposta no Portal 1"""
        try:
            # Encontrar e clicar no botão de envio
            botao_enviar = driver.find_element(By.ID, "submit-proposal")
            botao_enviar.click()
            
            # Esperar confirmação
            WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((By.CLASS_NAME, "success-message"))
            )
            
        except Exception as e:
            raise Exception(f"Erro ao enviar proposta no Portal 1: {str(e)}")
            
    def enviar_proposta_portal2(self, driver):
        """Envia a proposta no Portal 2"""
        try:
            # Encontrar e clicar no botão de envio
            botao_enviar = driver.find_element(By.ID, "btnSubmit")
            botao_enviar.click()
            
            # Esperar confirmação
            WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((By.CLASS_NAME, "proposal-sent"))
            )
            
        except Exception as e:
            raise Exception(f"Erro ao enviar proposta no Portal 2: {str(e)}")
            
    def finalizar_processamento(self, mensagem):
        self.processando = False
        self.parar_processamento = False
        self.btn_iniciar.config(state=tk.NORMAL)
        self.btn_parar.config(state=tk.DISABLED)
        self.status_label.config(text="Status: Aguardando")
        self.adicionar_log(mensagem)
        self.adicionar_output(f"\n{mensagem}", "success")
    
    def iniciar_processamento(self):
        if self.processando:
            return
            
        self.processando = True
        self.parar_processamento = False
        self.btn_iniciar.config(state=tk.DISABLED)
        self.btn_parar.config(state=tk.NORMAL)
        self.status_label.config(text="Status: Processando")
        
        # Limpar área de output
        self.output_area.delete('1.0', tk.END)
        
        # Executar script de informações gerais
        if self.executar_script_info_gerais():
            self.adicionar_log("Script executado, aguardando 2 segundos...")
            self.root.after(2000, self.carregar_dados)
        else:
            self.finalizar_processamento("Erro ao iniciar processamento")
            
    def carregar_dados(self):
        self.adicionar_log("Carregando dados do arquivo JSON...")
        self.dados_json = self.carregar_dados_json()
        
        if self.dados_json:
            self.adicionar_log("Dados JSON carregados com sucesso")
            self.carregar_cards_aceitos()
            
            if self.cards_data:
                self.adicionar_log(f"Iniciando processamento com {len(self.cards_data)} cards...")
                self.root.after(1000, self.processar_proximo_item)
            else:
                self.finalizar_processamento("Nenhum card em estado 'ACEITAS' encontrado")
        else:
            self.finalizar_processamento("Erro ao carregar dados JSON")

    def carregar_dados_json(self):
        try:
            # Usar caminho absoluto
            diretorio_base = os.path.dirname(os.path.abspath(__file__))
            arquivo_json = os.path.join(diretorio_base, 'json', 'dados_operacao.json')
            
            self.adicionar_log(f"Tentando carregar arquivo: {arquivo_json}")
            
            with open(arquivo_json, 'r', encoding='utf-8') as f:
                dados = json.load(f)
                total_empresas = len(dados)
                self.adicionar_log(f"Arquivo JSON carregado com sucesso. Total de empresas: {total_empresas}")
                return dados
        except Exception as e:
            self.adicionar_log(f"Erro ao carregar arquivo JSON: {str(e)}")
            return None

    def carregar_cards_aceitos(self):
        if not self.dados_json:
            self.adicionar_log("Nenhum dado JSON carregado")
            return
        
        self.cards_data = []
        for empresa_id, empresa_data in self.dados_json.items():
            etapas = empresa_data.get('etapas', {})
            if 'ACEITAS' in etapas:
                cards_aceitos = etapas['ACEITAS'].get('cards', [])
                for card in cards_aceitos:
                    card_processado = {
                        'id': card.get('ID', 'N/A'),
                        'pregao': card.get('Número do pregão', 'N/A'),
                        'portal': card.get('Portal', 'N/A'),
                        'empresa': empresa_data.get('nome_empresa', 'N/A'),
                        'empresa_id': empresa_id,
                        'status': 'ACEITAS',
                        'dados_completos': card
                    }
                    self.cards_data.append(card_processado)
        
        self.atualizar_lista_cards()
        self.atualizar_contador()

    def atualizar_lista_cards(self):
        # Limpar lista atual
        self.listbox.delete(0, tk.END)
        
        # Adicionar cards
        for card in self.cards_data:
            # Formatação mais informativa
            texto = f"Pregão: {card['pregao']} | Empresa: {card['empresa']} | Portal: {card['portal']}"
            self.listbox.insert(tk.END, texto)
    
    def processar_proximo_item(self):
        if self.parar_processamento:
            self.finalizar_processamento("Processamento interrompido pelo usuário")
            return
            
        if not self.cards_data:
            self.finalizar_processamento("Fila vazia - processamento concluído")
            return
            
        # Pegar o primeiro item da fila
        item = self.cards_data[0]
        self.cards_data = self.cards_data[1:]  # Remove o primeiro item
        
        # Atualizar lista
        self.atualizar_lista_cards()
        self.atualizar_contador()
        
        # Debug dos dados do item
        self.adicionar_output("\nDados do item:", "info")
        for key, value in item.items():
            self.adicionar_output(f"{key}: {value}", "info")
        
        # Identificar portal e processar
        portal = item.get('portal', '').lower()
        try:
            resultado = None
            if 'portal1' in portal:
                resultado = self.processar_portal1(item)
            elif 'portal2' in portal:
                resultado = self.processar_portal2(item)
            else:
                raise Exception(f"Portal não reconhecido: {portal}")
            
            # Adicionar aos resultados
            texto = f"Pregão: {item['pregao']} | Empresa: {item['empresa']} | Portal: {item['portal']}"
            if resultado:
                self.adicionar_resultado(texto, True)
            else:
                self.adicionar_resultado(texto, False)
                
        except Exception as e:
            texto = f"ERRO - Pregão: {item['pregao']} | Empresa: {item['empresa']} | Erro: {str(e)}"
            self.adicionar_resultado(texto, False)
        
        # Agendar próximo item
        self.root.after(1000, self.processar_proximo_item)
    
    def executar_script_info_gerais(self):
        self.adicionar_log("Executando script de coleta de informações gerais...")
        script_path = os.path.join("bitrix", "pega_informacoes_gerais.py")
        
        try:
            # Limpar área de output
            self.output_area.delete('1.0', tk.END)
            self.adicionar_output("Iniciando execução do script...\n", "info")
            
            # Executar o script e capturar a saída em tempo real
            process = subprocess.Popen(
                ["python", script_path],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                bufsize=1,
                universal_newlines=True
            )
            
            # Função para ler a saída em tempo real
            def read_output():
                # Ler stdout
                stdout_line = process.stdout.readline()
                if stdout_line:
                    self.adicionar_output(stdout_line.rstrip(), "info")
                    self.root.after(10, read_output)
                    return
                
                # Ler stderr
                stderr_line = process.stderr.readline()
                if stderr_line:
                    self.adicionar_output(stderr_line.rstrip(), "error")
                    self.root.after(10, read_output)
                    return
                
                # Se não há mais saída, verificar se o processo terminou
                if process.poll() is not None:
                    # Processo terminou
                    if process.returncode == 0:
                        self.adicionar_output("\nScript executado com sucesso!", "success")
                        self.root.after(2000, self.carregar_dados)
                    else:
                        self.adicionar_output("\nErro ao executar o script!", "error")
                        self.btn_iniciar.config(state=tk.NORMAL)
                    return
                
                # Se o processo ainda está rodando, continuar lendo
                self.root.after(10, read_output)
            
            # Iniciar a leitura da saída
            self.root.after(10, read_output)
            return True
            
        except Exception as e:
            self.adicionar_log(f"Erro ao executar o script: {str(e)}")
            self.adicionar_output(f"Erro: {str(e)}", "error")
            return False

    def adicionar_output(self, texto, tag=None):
        """Adiciona texto à área de output com a tag especificada"""
        if hasattr(self, 'output_area'):
            self.output_area.insert(tk.END, str(texto) + "\n", tag)
            self.output_area.see(tk.END)
            self.output_area.update_idletasks()

    def criar_painel_fila(self):
        # Painel de Fila (Central)
        fila_frame = tk.LabelFrame(
            self.main_frame,
            text="Fila de Processos",
            padx=10, pady=10,
            bg=self.frame_bg,
            font=('Helvetica', 10, 'bold')
        )
        fila_frame.grid(row=0, column=1, sticky='nsew', padx=5)
        
        # Título e contador
        header_frame = tk.Frame(fila_frame, bg=self.frame_bg)
        header_frame.pack(fill='x', pady=(0, 10))
        
        tk.Label(
            header_frame,
            text="Processos Aceitos",
            font=('Helvetica', 14, 'bold'),
            bg=self.frame_bg
        ).pack(side=tk.LEFT, pady=10)
        
        self.contador_label = tk.Label(
            header_frame,
            text="Total: 0",
            font=('Helvetica', 10),
            bg=self.frame_bg,
            fg='#666666'
        )
        self.contador_label.pack(side=tk.RIGHT, pady=10)
        
        # Lista de cards com scrollbar
        cards_frame = tk.Frame(fila_frame, bg=self.frame_bg)
        cards_frame.pack(fill='both', expand=True)
        
        scrollbar = tk.Scrollbar(cards_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.listbox = tk.Listbox(
            cards_frame,
            yscrollcommand=scrollbar.set,
            font=('Helvetica', 10),
            bg=self.highlight_color,
            selectmode=tk.SINGLE,
            activestyle='none'
        )
        self.listbox.pack(fill='both', expand=True)
        scrollbar.config(command=self.listbox.yview)
        
        # Inicializar lista vazia
        self.cards_data = []
    
    def criar_painel_detalhes(self):
        # Painel de Detalhes (Direito)
        detalhes_frame = tk.LabelFrame(
            self.main_frame,
            text="Saída do Terminal",
            padx=10, pady=10,
            bg=self.frame_bg,
            font=('Helvetica', 10, 'bold')
        )
        detalhes_frame.grid(row=0, column=2, sticky='nsew', padx=5)
        
        # Título
        tk.Label(
            detalhes_frame,
            text="Log de Execução",
            font=('Helvetica', 14, 'bold'),
            bg=self.frame_bg
        ).pack(pady=10)
        
        # Frame para o texto
        text_frame = tk.Frame(detalhes_frame, bg=self.frame_bg)
        text_frame.pack(fill='both', expand=True)
        
        # Scrollbars
        vsb = ttk.Scrollbar(text_frame, orient="vertical")
        hsb = ttk.Scrollbar(text_frame, orient="horizontal")
        
        # Área de texto para output
        self.output_area = tk.Text(
            text_frame,
            wrap=tk.NONE,
            font=('Consolas', 10),
            bg='black',
            fg='white',
            yscrollcommand=vsb.set,
            xscrollcommand=hsb.set
        )
        
        # Configurar scrollbars
        vsb.config(command=self.output_area.yview)
        hsb.config(command=self.output_area.xview)
        
        # Posicionar elementos
        vsb.pack(side='right', fill='y')
        hsb.pack(side='bottom', fill='x')
        self.output_area.pack(fill='both', expand=True)
        
        # Configurar tags para cores
        self.output_area.tag_configure("error", foreground="red")
        self.output_area.tag_configure("success", foreground="green")
        self.output_area.tag_configure("info", foreground="cyan")
        
        # Mensagem inicial
        self.adicionar_output("Sistema iniciado. Aguardando processamento...", "info")
        
    def adicionar_log(self, mensagem):
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.log_area.insert(tk.END, f"[{timestamp}] {mensagem}\n")
        self.log_area.see(tk.END)
    
    def criar_painel_fila(self):
        # Frame principal para fila e resultados
        main_frame = tk.Frame(self.main_frame, bg=self.bg_color)
        main_frame.grid(row=0, column=1, sticky='nsew', padx=5)
        
        # Configurar grid do frame principal
        main_frame.grid_rowconfigure(0, weight=2)  # Fila ocupa 2/3
        main_frame.grid_rowconfigure(1, weight=1)  # Resultados ocupa 1/3
        main_frame.grid_columnconfigure(0, weight=1)
        
        # Painel de Fila (Superior)
        fila_frame = tk.LabelFrame(
            main_frame,
            text="Fila de Processos",
            padx=10, pady=10,
            bg=self.frame_bg,
            font=('Helvetica', 10, 'bold')
        )
        fila_frame.grid(row=0, column=0, sticky='nsew', padx=5, pady=(0,5))
        
        # Frame para o título e contador
        titulo_frame = tk.Frame(fila_frame, bg=self.frame_bg)
        titulo_frame.pack(fill='x', pady=(0, 10))
        
        # Contador
        self.contador_label = tk.Label(
            titulo_frame,
            text="Total: 0",
            font=('Helvetica', 10),
            bg=self.frame_bg
        )
        self.contador_label.pack(side='right', pady=5)
        
        # Frame para a lista
        list_frame = tk.Frame(fila_frame, bg=self.frame_bg)
        list_frame.pack(fill='both', expand=True)
        
        # Scrollbar
        scrollbar = tk.Scrollbar(list_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Listbox
        self.listbox = tk.Listbox(
            list_frame,
            yscrollcommand=scrollbar.set,
            font=('Helvetica', 10),
            bg=self.highlight_color,
            selectmode=tk.SINGLE,
            activestyle='none'
        )
        self.listbox.pack(fill='both', expand=True)
        scrollbar.config(command=self.listbox.yview)
        
        # Painel de Resultados (Inferior)
        resultados_frame = tk.LabelFrame(
            main_frame,
            text="Resultados do Processamento",
            padx=10, pady=10,
            bg=self.frame_bg,
            font=('Helvetica', 10, 'bold')
        )
        resultados_frame.grid(row=1, column=0, sticky='nsew', padx=5, pady=(5,0))
        
        # Criar notebook para separar sucesso e erro
        notebook = ttk.Notebook(resultados_frame)
        notebook.pack(fill='both', expand=True)
        
        # Tab de Sucesso
        sucesso_frame = tk.Frame(notebook, bg=self.frame_bg)
        notebook.add(sucesso_frame, text='Sucesso')
        
        # Lista de sucesso
        self.sucesso_listbox = tk.Listbox(
            sucesso_frame,
            font=('Helvetica', 10),
            bg=self.highlight_color,
            fg='green'
        )
        self.sucesso_listbox.pack(fill='both', expand=True)
        
        # Tab de Erro
        erro_frame = tk.Frame(notebook, bg=self.frame_bg)
        notebook.add(erro_frame, text='Erro')
        
        # Lista de erro
        self.erro_listbox = tk.Listbox(
            erro_frame,
            font=('Helvetica', 10),
            bg=self.highlight_color,
            fg='red'
        )
        self.erro_listbox.pack(fill='both', expand=True)

    def adicionar_resultado(self, texto, sucesso=True):
        """Adiciona um resultado na lista apropriada"""
        if sucesso:
            self.sucesso_listbox.insert(0, texto)  # Inserir no topo
        else:
            self.erro_listbox.insert(0, texto)  # Inserir no topo

    def carregar_dados(self):
        self.dados_json = self.carregar_dados_json()
        if self.dados_json:
            self.carregar_cards_aceitos()
            if self.cards_data:
                self.root.after(1000, self.processar_proximo_item)
            else:
                self.finalizar_processamento("Nenhum card em estado 'ACEITAS' encontrado")
        else:
            self.finalizar_processamento("Erro ao carregar dados JSON")

    def carregar_cards_aceitos(self):
        if not self.dados_json:
            return
        
        self.cards_data = []
        for empresa_id, empresa_data in self.dados_json.items():
            etapas = empresa_data.get('etapas', {})
            if 'ACEITAS' in etapas:
                cards_aceitos = etapas['ACEITAS'].get('cards', [])
                for card in cards_aceitos:
                    card_processado = {
                        'id': card.get('ID', 'N/A'),
                        'pregao': card.get('Número do pregão', 'N/A'),
                        'portal': card.get('Portal', 'N/A'),
                        'empresa': empresa_data.get('nome_empresa', 'N/A'),
                        'empresa_id': empresa_id,
                        'status': 'ACEITAS',
                        'dados_completos': card
                    }
                    self.cards_data.append(card_processado)
        
        self.atualizar_lista_cards()
        self.atualizar_contador()

    def processar_proximo_item(self):
        if self.parar_processamento:
            self.finalizar_processamento("Processamento interrompido pelo usuário")
            return
            
        if not self.cards_data:
            self.finalizar_processamento("Fila vazia - processamento concluído")
            return
            
        # Pegar o primeiro item da fila
        item = self.cards_data[0]
        self.cards_data = self.cards_data[1:]  # Remove o primeiro item
        
        # Atualizar lista
        self.atualizar_lista_cards()
        self.atualizar_contador()
        
        # Identificar portal e processar
        portal = item.get('portal', '').lower()
        try:
            resultado = None
            if 'portal1' in portal:
                resultado = self.processar_portal1(item)
            elif 'portal2' in portal:
                resultado = self.processar_portal2(item)
            else:
                raise Exception(f"Portal não reconhecido: {portal}")
            
            # Adicionar aos resultados
            texto = f"Pregão: {item['pregao']} | Empresa: {item['empresa']} | Portal: {item['portal']}"
            if resultado:
                self.adicionar_resultado(texto, True)
            else:
                self.adicionar_resultado(texto, False)
                
        except Exception as e:
            texto = f"ERRO - Pregão: {item['pregao']} | Empresa: {item['empresa']} | Erro: {str(e)}"
            self.adicionar_resultado(texto, False)
        
        # Agendar próximo item
        self.root.after(1000, self.processar_proximo_item)
    
    def executar_script_info_gerais(self):
        self.adicionar_log("Executando script de coleta de informações gerais...")
        script_path = os.path.join("bitrix", "pega_informacoes_gerais.py")
        
        try:
            # Limpar área de output
            self.output_area.delete('1.0', tk.END)
            self.adicionar_output("Iniciando execução do script...\n", "info")
            
            # Executar o script e capturar a saída em tempo real
            process = subprocess.Popen(
                ["python", script_path],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                bufsize=1,
                universal_newlines=True
            )
            
            # Função para ler a saída em tempo real
            def read_output():
                # Ler stdout
                stdout_line = process.stdout.readline()
                if stdout_line:
                    self.adicionar_output(stdout_line.rstrip(), "info")
                    self.root.after(10, read_output)
                    return
                
                # Ler stderr
                stderr_line = process.stderr.readline()
                if stderr_line:
                    self.adicionar_output(stderr_line.rstrip(), "error")
                    self.root.after(10, read_output)
                    return
                
                # Se não há mais saída, verificar se o processo terminou
                if process.poll() is not None:
                    # Processo terminou
                    if process.returncode == 0:
                        self.adicionar_output("\nScript executado com sucesso!", "success")
                        self.root.after(2000, self.carregar_dados)
                    else:
                        self.adicionar_output("\nErro ao executar o script!", "error")
                        self.btn_iniciar.config(state=tk.NORMAL)
                    return
                
                # Se o processo ainda está rodando, continuar lendo
                self.root.after(10, read_output)
            
            # Iniciar a leitura da saída
            self.root.after(10, read_output)
            return True
            
        except Exception as e:
            self.adicionar_log(f"Erro ao executar o script: {str(e)}")
            self.adicionar_output(f"Erro: {str(e)}", "error")
            return False

    def atualizar_contador(self):
        """Atualiza o contador de itens na fila"""
        if hasattr(self, 'contador_label'):
            total = len(self.cards_data) if hasattr(self, 'cards_data') else 0
            self.contador_label.config(text=f"Total: {total}")

if __name__ == "__main__":
    app = TelaPrincipal()
    app.root.mainloop()