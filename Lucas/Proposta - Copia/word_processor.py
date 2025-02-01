from docx import Document
from docx.shared import Pt, Cm, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
import json
import os

def configurar_pagina_a4(doc):
    """
    Configura o documento para usar o formato A4
    """
    section = doc.sections[0]
    section.page_height = Mm(297)  # A4 altura
    section.page_width = Mm(210)   # A4 largura
    # Margens
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    return section

def calcular_larguras_colunas():
    """
    Calcula as larguras das colunas baseado em proporções fixas
    Retorna larguras em centímetros
    """
    # Largura total disponível em uma página A4 com margens de 1 polegada
    # A4 = 21cm, menos 2 polegadas de margem (5.08cm) = ~15.92cm
    largura_disponivel = 15.92
    
    # Define as proporções de cada coluna (total = 100)
    proporcoes = [
        7,    # Item (7%)
        30,   # Descrição (30%)
        13,   # Marca/Fabricante (13%)
        13,   # Modelo/Versão (13%)
        9,    # Quantidade (9%)
        9,    # Unidade (9%)
        9,    # Valor Unitário (9%)
        10    # Valor Total (10%)
    ]
    
    # Calcula as larguras em centímetros
    return [largura_disponivel * prop / 100 for prop in proporcoes]

def calcular_valor_unitario(valor_total, quantidade):
    """
    Calcula o valor unitário baseado no valor total e quantidade
    """
    try:
        valor_total = float(str(valor_total).replace('R$', '').replace('.', '').replace(',', '.').strip())
        quantidade = float(str(quantidade).strip())
        if quantidade > 0:
            return f"R$ {(valor_total / quantidade):,.2f}".replace(',', '_').replace('.', ',').replace('_', '.')
    except:
        pass
    return ''

def criar_tabela_word(json_file, template_file=None, marcador_tabela="{{TABELA_AQUI}}", output_file=None):
    """
    Cria um documento Word com uma tabela baseada nos dados do JSON.
    Se um template for fornecido, procura pelo marcador e insere a tabela naquela posição.
    """
    try:
        # Lê o arquivo JSON
        with open(json_file, 'r', encoding='utf-8') as f:
            dados = json.load(f)

        # Cria ou abre o documento Word
        if template_file and os.path.exists(template_file):
            doc = Document(template_file)
            print(f"Usando template: {template_file}")
            tem_template = True
        else:
            doc = Document()
            tem_template = False
            print("Criando novo documento")
            
        # Configura para A4
        configurar_pagina_a4(doc)
        
        # Calcula as larguras das colunas
        larguras = calcular_larguras_colunas()

        # Define o mapeamento de colunas (cabeçalho -> nome no JSON)
        mapeamento_colunas = {
            "Item": "Item",
            "Descrição": "Descrição Detalhada do Objeto Ofertado",
            "Marca/Fabricante": "Marca/Fabricante",
            "Modelo/Versão": "Modelo/Versão",
            "Quantidade": "Quantidade solicitada",
            "Unidade": "Unidade de Fornecimento",
            "Valor Unitário": None,  # Será calculado
            "Valor Total": "Valor Total"
        }

        # Cria a tabela temporária
        table = doc.add_table(rows=1, cols=len(mapeamento_colunas))
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        
        # Preenche o cabeçalho
        header_cells = table.rows[0].cells
        for i, (coluna_exibicao, _) in enumerate(mapeamento_colunas.items()):
            header_cells[i].text = coluna_exibicao
            header_cells[i].width = Cm(larguras[i])
            paragraph = header_cells[i].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0]
            run.font.bold = True
            run.font.size = Pt(10)

        # Preenche os dados
        for item in dados:
            row_cells = table.add_row().cells
            for i, (coluna_exibicao, coluna_json) in enumerate(mapeamento_colunas.items()):
                if coluna_exibicao == "Valor Unitário":
                    # Calcula o valor unitário
                    valor = calcular_valor_unitario(
                        item.get("Valor Total", ""),
                        item.get("Quantidade solicitada", "")
                    )
                else:
                    valor = str(item.get(coluna_json, ""))
                
                row_cells[i].text = valor
                row_cells[i].width = Cm(larguras[i])
                paragraph = row_cells[i].paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if i == 1:  # Coluna de descrição
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = paragraph.runs[0]
                run.font.size = Pt(9)

        # Se tiver template, procura o marcador e move a tabela
        if tem_template:
            # Procura o parágrafo com o marcador
            marcador_encontrado = False
            tabela_elemento = table._element
            
            # Remove a tabela temporária do documento
            tabela_elemento.getparent().remove(tabela_elemento)
            
            for i, paragraph in enumerate(doc.paragraphs):
                if marcador_tabela in paragraph.text:
                    print(f"Marcador '{marcador_tabela}' encontrado no documento")
                    # Remove o parágrafo com o marcador
                    p = paragraph._element
                    p.getparent().remove(p)
                    # Insere a tabela no local exato onde estava o marcador
                    doc._body._body.insert(i, tabela_elemento)
                    marcador_encontrado = True
                    break
            
            if not marcador_encontrado:
                print(f"AVISO: Marcador '{marcador_tabela}' não encontrado no documento!")
                # Se não encontrou o marcador, adiciona a tabela no final
                print("Adicionando tabela no final do documento")
                doc._body._body.append(tabela_elemento)

        # Tenta salvar o arquivo com o nome especificado
        if output_file:
            try:
                doc.save(output_file)
                print(f"\nDocumento Word salvo com sucesso como: {output_file}")
                return output_file
            except PermissionError:
                # Se der erro de permissão, tenta salvar com um nome alternativo
                dir_name = os.path.dirname(output_file)
                base_name = os.path.basename(output_file)
                nome, ext = os.path.splitext(base_name)
                novo_arquivo = os.path.join(dir_name, f"{nome}_novo{ext}")
                print(f"Arquivo original está em uso. Salvando como: {novo_arquivo}")
                doc.save(novo_arquivo)
                return novo_arquivo
        else:
            # Se não foi especificado um nome, usa um padrão
            output_file = os.path.join(os.path.dirname(json_file), 'proposta_precos.docx')
            doc.save(output_file)
            return output_file

    except Exception as e:
        print(f"Erro ao criar documento Word: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return None
