import os
import json
import time
import subprocess
from datetime import datetime
import pandas as pd
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import Select
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
import tkinter as tk
from tkinter import ttk, scrolledtext
import glob
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fila_processor import FilaProcessor
from portal_processor import PortalProcessor

# Dicionário com as modalidades por portal
MODALIDADES = {
    'bll': ['Selecione a modalidade', 'Pregão', 'Concorrência', 'Dispensa'],
    'comprasnet': ['Selecione a modalidade', 'Pregão', 'Concorrência', 'Dispensa'],
    'compraspublicas': ['Selecione a modalidade', 'Pregão', 'Concorrência', 'Dispensa'],
    'bnc': ['Selecione a modalidade', 'Pregão', 'Concorrência', 'Dispensa'],
    'comprasbr': ['Selecione a modalidade', 'Pregão', 'Concorrência', 'Dispensa'],
    'licitanet': ['Selecione a modalidade', 'Pregão', 'Concorrência', 'Dispensa']
}

# Mapeamento de valores das modalidades por portal
MODALIDADE_VALUES = {
    'bll': {'Selecione a modalidade': '0', 'Pregão': '1', 'Concorrência': '3', 'Dispensa': '10'},
    'comprasnet': {'Selecione a modalidade': '0', 'Pregão': '1', 'Concorrência': '3', 'Dispensa': '10'},
    'compraspublicas': {'Selecione a modalidade': '0', 'Pregão': '1', 'Concorrência': '3', 'Dispensa': '10'},
    'bnc': {'Selecione a modalidade': '0', 'Pregão': '1', 'Concorrência': '3', 'Dispensa': '10'},
    'comprasbr': {'Selecione a modalidade': '0', 'Pregão': '1', 'Concorrência': '3', 'Dispensa': '10'},
    'licitanet': {'Selecione a modalidade': '0', 'Pregão': '1', 'Concorrência': '3', 'Dispensa': '10'}
}

def listar_portais(self):         
    # Lista de portais
    portais = [
        ("BLL - Bolsa de Licitações do Brasil", "bll"),
        ("ComprasNet - Compras Governamentais", "comprasnet"),
        ("Portal de Compras Públicas", "compraspublicas"),
        ("BNC - Bolsa Nacional de Compras", "bnc"),
        ("Compras BR - Portal de Compras do Brasil", "comprasbr"),
        ("LicitaNet - Portal de Licitações", "licitanet")
    ]

    def abrir_firefox(self):
        options = Options()
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--disable-extensions")
        options.add_argument("--disable-gpu")
        options.add_argument("--disable-notifications")
        options.add_argument("--headless")
    
    
    def encontrar_ultimo_excel(self):
        """Encontra o arquivo Excel mais recente na pasta de downloads."""
        downloads_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'downloads')
        # Procura por arquivos que começam com 'exportacao_' e terminam com .xlsx
        lista_arquivos = glob.glob(os.path.join(downloads_path, 'exportacao_*.xlsx'))
        if not lista_arquivos:
            return None
        # Retorna o arquivo mais recente
        return max(lista_arquivos, key=os.path.getctime)

    def processar_excel(self, arquivo_excel):
        """Processa o arquivo Excel, modificando valores específicos."""
        try:
            # Lê o arquivo original com openpyxl engine
            print("Iniciando leitura do arquivo Excel...")
            df = pd.read_excel(arquivo_excel, engine='openpyxl')
            new_df = df.copy()

            print(f"Arquivo lido com sucesso. Colunas encontradas: {list(df.columns)}")

            # Modifica valores da terceira coluna
            if len(df.columns) > 2:
                third_column = df.columns[2]
                new_df.iloc[1:, 2] = 0
                print(f"Valores da terceira coluna ({third_column}) modificados")

            # Preenche campos "marca" e "modelo"
            colunas_lower = [col.lower() for col in new_df.columns]
            if 'marca' in colunas_lower:
                idx = colunas_lower.index('marca')
                new_df.iloc[:, idx] = 'Própria'
                print("Campo 'marca' preenchido")
            if 'modelo' in colunas_lower:
                idx = colunas_lower.index('modelo')
                new_df.iloc[:, idx] = 'Própria'
                print("Campo 'modelo' preenchido")

            # Cria o nome do arquivo modificado
            nome_base = os.path.splitext(os.path.basename(arquivo_excel))[0]
            novo_arquivo = arquivo_excel.replace('.xlsx', '_modificado.xlsx')
            
            # Salva o arquivo modificado
            print(f"Salvando arquivo modificado como: {novo_arquivo}")
            new_df.to_excel(novo_arquivo, index=False, engine='openpyxl')
            print(f"Excel processado e salvo com sucesso")
            return novo_arquivo
        except Exception as e:
            print(f"Erro ao processar Excel: {str(e)}")
            print("Detalhes do erro:")
            import traceback
            print(traceback.format_exc())
            return None

    def excel_para_json(self, arquivo_excel):
        """Converte o arquivo Excel para JSON com formato específico."""
        try:
            # Lê o arquivo Excel com openpyxl engine
            print("Iniciando conversão do Excel para JSON...")
            df = pd.read_excel(arquivo_excel, engine='openpyxl')
            
            # Prepara os dados no formato desejado
            nome_base = os.path.splitext(os.path.basename(arquivo_excel))[0]
            data = {
                "processo": nome_base,
                "colunas": list(df.columns),
                "dados": df.to_dict(orient="records")
            }
            
            # Cria o nome do arquivo JSON
            arquivo_json = arquivo_excel.replace('.xlsx', '.json')
            
            # Salva o JSON
            print(f"Salvando dados em: {arquivo_json}")
            with open(arquivo_json, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=4)
            
            print(f"Arquivo JSON criado com sucesso")
            return arquivo_json
        except Exception as e:
            print(f"Erro ao converter Excel para JSON: {str(e)}")
            print("Detalhes do erro:")
            import traceback
            print(traceback.format_exc())
            return None

    def excluir_excel_temporario(self, arquivo_excel):
        """Exclui o arquivo Excel após a conversão para JSON."""
        try:
            if os.path.exists(arquivo_excel):
                os.remove(arquivo_excel)
                print(f"Arquivo Excel excluído: {arquivo_excel}")
        except Exception as e:
            print(f"Erro ao excluir arquivo Excel: {str(e)}")

    def processar_arquivo_exportado(self, arquivo_excel):
        """
        Processa o arquivo Excel exportado e prepara os dados para o Word
        """
        try:
            print(f"\nProcessando arquivo Excel: {arquivo_excel}")
            # Lê o arquivo Excel
            df = pd.read_excel(arquivo_excel, engine='openpyxl')
            
            # Remove linhas vazias
            df = df.dropna(how='all')
            
            # Processa as colunas necessárias
            dados_processados = []
            for index, row in df.iterrows():
                item = {
                    'item': row.get('Item', ''),
                    'descricao': row.get('Descrição', ''),
                    'unidade': row.get('Unidade de fornecimento', ''),
                    'quantidade': row.get('Quantidade', 0),
                    'valor_unitario': row.get('Valor Unitário', 0),
                    'valor_total': row.get('Valor Total', 0)
                }
                dados_processados.append(item)
            
            print(f"Processados {len(dados_processados)} itens do Excel")
            return dados_processados
            
        except Exception as e:
            print(f"Erro ao processar arquivo Excel: {str(e)}")
            print("Detalhes do erro:")
            import traceback
            print(traceback.format_exc())
            return None

    def gerar_documento_word(self, dados_processados, template_word=None):
        """
        Gera um documento Word com os dados processados
        """
        try:
            from docx import Document
            from docx.shared import Pt
            
            # Se não fornecido template, cria novo documento
            if template_word and os.path.exists(template_word):
                doc = Document(template_word)
                print(f"Usando template: {template_word}")
            else:
                doc = Document()
                print("Criando novo documento Word")
            
            # Adiciona tabela com os dados
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            
            # Cabeçalho
            header_cells = table.rows[0].cells
            headers = ['Item', 'Descrição', 'Unidade', 'Quantidade', 'Valor Unitário', 'Valor Total']
            for i, header in enumerate(headers):
                header_cells[i].text = header
            
            # Dados
            for item in dados_processados:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item['item'])
                row_cells[1].text = str(item['descricao'])
                row_cells[2].text = str(item['unidade'])
                row_cells[3].text = str(item['quantidade'])
                row_cells[4].text = f"R$ {item['valor_unitario']:.2f}"
                row_cells[5].text = f"R$ {item['valor_total']:.2f}"
            
            # Salva o documento
            output_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'proposta_processada.docx')
            doc.save(output_file)
            print(f"\nDocumento Word gerado com sucesso: {output_file}")
            return output_file
            
        except Exception as e:
            print(f"Erro ao gerar documento Word: {str(e)}")
            print("Detalhes do erro:")
            import traceback
            print(traceback.format_exc())
            return None

    def processar_arquivo(self, arquivo_excel, arquivo_word):
        """
        Processa o arquivo Excel e gera o documento Word
        """
        try:
            print("\nIniciando processamento do arquivo...")
            
            # Processa o Excel e gera o JSON
            resultado = self.processar_arquivo_exportado(arquivo_excel)
            if not resultado:
                print("Erro ao processar o arquivo Excel")
                return False
            
            # Remove o arquivo Excel processado pois não será mais necessário
            try:
                os.remove(resultado['excel'])
                print(f"\nArquivo Excel original removido: {resultado['excel']}")
            except:
                pass
            
            # Lista de marcadores para tentar
            marcadores = [
                "{{TABELA_AQUI}}",
                "{{tabela_aqui}}",
                "{TABELA_AQUI}",
                "{tabela_aqui}",
                "TABELA_AQUI",
                "tabela_aqui"
            ]
            
            # Tenta gerar o documento Word com cada marcador
            for marcador in marcadores:
                print(f"\nTentando com marcador: {marcador}")
                
                # Define o nome do arquivo de saída baseado no template
                template_dir = os.path.dirname(arquivo_word)
                template_nome = os.path.basename(arquivo_word)
                nome_base = os.path.splitext(template_nome)[0]
                arquivo_saida = os.path.join(template_dir, f"{nome_base}_proposta.docx")
                
                # Gera o documento Word
                doc_gerado = criar_tabela_word(
                    resultado['json'],
                    template_file=arquivo_word,
                    marcador_tabela=marcador,
                    output_file=arquivo_saida
                )
                
                if doc_gerado:
                    print(f"Documento Word gerado com sucesso usando marcador: {marcador}")
                    print(f"Documento Word gerado: {doc_gerado}")
                    return True
                
            print("\nNão foi possível gerar o documento Word com nenhum dos marcadores tentados")
            return False
        
        except Exception as e:
            print(f"Erro durante o processamento: {str(e)}")
            return False
        finally:
            print("\nProcessamento concluído!")

    def voltar(self):
        self.root.destroy()
        SeletorPortal()

if __name__ == "__main__":
    app = TelaPrincipal()
    app.root.mainloop()